import sqlite3
from flask import Flask, render_template, redirect, request, flash, send_file, send_from_directory, url_for
from werkzeug.exceptions import abort
import os
from docx import Document
from init_db import initialize_database

if __name__ == '__main__':
    initialize_database()  
    
app = Flask(__name__)
app.config['SECRET_KEY'] = 'my)secret)key'
UPLOAD_FOLDER = 'contracts'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def get_db_connection():
    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    return conn

@app.route('/')
def index():
    return redirect("/contracts")

# КОНТРАКТЫ  

@app.route('/contracts')
def contracts():
    conn = get_db_connection()
    contracts_data = conn.execute("""
        SELECT 
            contracts.*,
            clients.name AS client_name,
            services.name AS service_name,
            employees.name AS employee_name
        FROM contracts
        JOIN clients ON contracts.client_id = clients.id_client
        JOIN services ON contracts.service_id = services.service_id
        JOIN employees ON contracts.employee_id = employees.id_employee;
    """).fetchall()
    conn.close()
    return render_template('contracts.html', contracts=contracts_data)

def get_contract(item_id):
    conn = get_db_connection()
    item = conn.execute("""
        SELECT 
            contracts.*,
            clients.name AS client_name,
            services.name AS service_name,
            employees.name AS employee_name
        FROM contracts
        JOIN clients ON contracts.client_id = clients.id_client
        JOIN services ON contracts.service_id = services.service_id
        JOIN employees ON contracts.employee_id = employees.id_employee
        WHERE contracts.id_contract = ?;
    """, (item_id,)).fetchone()
    conn.close()
    if item is None:
        abort(404)  
    return item

@app.route('/contract/<int:contract_id>')
def contract(contract_id):
    pos = get_contract(contract_id)
    return render_template('contract.html', contract=pos)

@app.route('/new_contract', methods=('GET', 'POST'))
def new_contract():
    if request.method == 'POST':
        try:
            number = request.form['number']
            date = request.form['date']
            start_price = int(request.form['start_price'])
            discount = int(request.form['discount'])
            deal_type = request.form['deal_type']
            finish_price = request.form['finish_price']
            client_id = int(request.form.get('client'))
            service_id = int(request.form.get('service'))
            employee_id = int(request.form.get('employee'))
            deal_status = 0
        except ValueError:
            flash('Некорректные значения')
            return render_template('new_contract.html')

        if not (client_id > 0 and service_id > 0 and employee_id > 0):
            flash('Не все поля заполнены')
        else:
            conn = get_db_connection()
            conn.execute("""
                INSERT INTO contracts 
                (number, date, deal_type, start_price, discount, deal_status, finish_price, client_id, service_id, employee_id) 
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (number, date, deal_type, start_price, discount, deal_status, finish_price, client_id, service_id, employee_id))
            conn.commit()
            conn.close()
            return redirect('/contracts')

    conn = get_db_connection()
    clients = conn.execute("SELECT * FROM clients").fetchall()
    services = conn.execute("SELECT * FROM services").fetchall()
    employees = conn.execute("SELECT * FROM employees").fetchall()
    conn.close()

    return render_template('new_contract.html', 
                           clients=clients, 
                           services=services,
                           employees=employees)

@app.route('/generate_contract', methods=('GET', 'POST'))
def generate_contract():
    id_contract = request.args.get('id_contract', 0, type=int)

    # Получение данных о договоре из базы
    conn = get_db_connection()
    contract = conn.execute("""
    SELECT 
        contracts.number AS CONTRACT_NUMBER,
        contracts.date AS CONTRACT_DATE,
        clients.name AS CLIENT_COMPANY_NAME,
        services.name AS SERVICE_NAME,
        services.description AS SERVICE_DESCRIPTION,
        contracts.finish_price AS PRICE,
        contracts.discount AS SALE,
        employees.name AS EMPLOYEE_FULLNAME
    FROM contracts
    JOIN clients ON contracts.client_id = clients.id_client
    JOIN services ON contracts.service_id = services.service_id
    JOIN employees ON contracts.employee_id = employees.id_employee
    WHERE contracts.id_contract = ?
""", (id_contract,)).fetchone()

    if not contract:
        flash("Договор не найден.")
        return redirect('/contracts')

    auto_params = {
    'CONTRACT_NUMBER': ['Номер договора', contract['CONTRACT_NUMBER']],
    'CONTRACT_DATE': ['Дата договора', contract['CONTRACT_DATE']],
    'CLIENT_COMPANY_NAME': ['Название клиента', contract['CLIENT_COMPANY_NAME']],
    'SERVICE_NAME': ['Название продукта', contract['SERVICE_NAME']],
    'SERVICE_DESCRIPTION': ['Описание услуги', contract['SERVICE_DESCRIPTION']],
    'PRICE': ['Стоимость услуги', f"{contract['PRICE']} рублей"],
    'SALE': ['Скидка', f"{contract['SALE']}%"],
    'COMPANY_NAME': ['Название компании', '"FOUR SEASONS MOSCOW"'],
    'EMPLOYEE_FULLNAME': ['ФИО сотрудника', contract['EMPLOYEE_FULLNAME']],
}


    if request.method == 'POST':
        template_path = os.path.join('templates', 'contract_template.docx')
        output_path = os.path.join('contracts', f"Договор_{auto_params['CONTRACT_NUMBER'][1]}.docx")
        generate_docx(template_path, output_path, auto_params)

        return send_file(output_path, as_attachment=True)
    return render_template(
        'generate_contract.html',
        auto_params=auto_params,
        contract_params={} 
    )

def generate_docx(template_path, output_path, params):
    """Генерация .docx файла на основе шаблона и параметров"""
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        for key, value in params.items():
            replacement = value[1] if isinstance(value, list) and len(value) > 1 else str(value)
            if f"=={key}==" in paragraph.text:
                paragraph.text = paragraph.text.replace(f"=={key}==", replacement)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in params.items():
                    replacement = value[1] if isinstance(value, list) and len(value) > 1 else str(value)
                    replacement = replacement.strip("[]")
                    if f"=={key}==" in cell.text:
                        cell.text = cell.text.replace(f"=={key}==", replacement)
    doc.save(output_path)

@app.route('/edit_contract/<int:contract_id>', methods=['GET', 'POST'])
def edit_contract(contract_id):
    conn = get_db_connection()
    contract = conn.execute('SELECT * FROM contracts WHERE id_contract = ?', (contract_id,)).fetchone()
    clients = conn.execute('SELECT * FROM clients').fetchall()
    services = conn.execute('SELECT * FROM services').fetchall()
    employees = conn.execute('SELECT * FROM employees').fetchall()
    if request.method == 'POST':
        number = request.form['number']
        date = request.form['date']
        deal_type = request.form['deal_type']
        start_price = request.form['start_price']
        discount = request.form['discount']
        finish_price = request.form['finish_price']
        client_id = request.form['client_id']
        service_id = request.form['service_id']
        employee_id = request.form['employee_id']
        conn.execute('''UPDATE contracts SET number = ?, date = ?, deal_type = ?, start_price = ?, discount = ?, finish_price = ?, client_id = ?, service_id = ?, employee_id = ? WHERE id_contract = ?''',
                     (number, date, deal_type, start_price, discount, finish_price, client_id, service_id, employee_id, contract_id))
        conn.commit()
        conn.close()
        flash('Контракт успешно обновлен.')
        return redirect('/contracts')
    conn.close()
    return render_template('edit_contract.html', contract=contract, clients=clients, services=services, employees=employees)

# КЛИЕНТЫ

@app.route('/clients')
def clients():
    conn = get_db_connection()
    pos = conn.execute("SELECT * FROM clients").fetchall()
    conn.close()
    return render_template('clients.html', clients=pos)

def get_client(item_id):
    conn = get_db_connection()
    client = conn.execute("SELECT * FROM clients WHERE id_client = ?", (item_id,)).fetchone()
    if not client:
        conn.close()
        abort(404) 
    contracts = conn.execute("""
        SELECT 
            contracts.id_contract, 
            contracts.number, 
            contracts.date,
            contracts.deal_type AS deal_type,
            contracts.start_price, 
            contracts.finish_price, 
            employees.name AS employee_name,
            services.name AS service_name
        FROM contracts
        JOIN employees ON contracts.employee_id = employees.id_employee
        JOIN services ON contracts.service_id = services.service_id
        WHERE contracts.client_id = ?
    """, (item_id,)).fetchall()
    conn.close()
    return {
        "client": client,
        "contracts": contracts
    }

@app.route('/client/<int:client_id>')
def client(client_id):
    data = get_client(client_id)
    return render_template('client.html', client=data["client"], contracts=data["contracts"])

@app.route('/edit_client/<int:client_id>', methods=['GET', 'POST'])
def edit_client(client_id):
    conn = get_db_connection()
    client = conn.execute('SELECT * FROM clients WHERE id_client = ?', (client_id,)).fetchone()
    if request.method == 'POST':
        name = request.form['name']
        email = request.form['email']
        contact = request.form['contact']
        address = request.form['address']
        conn.execute('''UPDATE clients SET name = ?, email = ?, contact = ?, address = ? WHERE id_client = ?''',
                     (name, email, contact, address, client_id))
        conn.commit()
        conn.close()
        flash('Клиент успешно обновлен.')
        return redirect('/clients')
    conn.close()
    return render_template('edit_client.html', client=client)

@app.route('/new_client', methods=['GET', 'POST'])
def new_client():
    if request.method == 'POST':
        try:
            name = request.form['name'].strip()
            email = request.form['email'].strip()
            contact = request.form['contact'].strip()
            address = request.form['address'].strip()
            if not (name and email and contact and address):
                flash("Все поля должны быть заполнены!", "danger")
                return render_template('new_client.html')
            conn = get_db_connection()
            conn.execute("""
                INSERT INTO clients (name, email, contact, address)
                VALUES (?, ?, ?, ?)
            """, (name, email, contact, address))
            conn.commit()
            conn.close()
            flash("Клиент успешно добавлен!", "success")
            return redirect('/clients')
        except Exception as e:
            flash(f"Ошибка при добавлении клиента: {e}", "danger")
            return render_template('new_client.html')
    return render_template('new_client.html')

 
@app.route('/services')
def services():
    conn = get_db_connection()
    pos = conn.execute("SELECT * FROM services").fetchall()
    conn.close()
    return render_template('services.html', services=pos)

def get_service(item_id):
    conn = get_db_connection()
    service = conn.execute("SELECT * FROM services WHERE service_id = ?", (item_id,)).fetchone()
    if not service:
        conn.close()
        abort(404)  
    contracts = conn.execute("""
        SELECT 
            contracts.id_contract,
            contracts.number,
            contracts.date,
            contracts.deal_type,
            contracts.start_price,
            contracts.finish_price,
            clients.name AS client_name,
            employees.name AS employee_name
        FROM contracts
        JOIN clients ON contracts.client_id = clients.id_client
        JOIN employees ON contracts.employee_id = employees.id_employee
        WHERE contracts.service_id = ?
    """, (item_id,)).fetchall()
    conn.close()
    return {
        "service": service,
        "contracts": contracts
    }

@app.route('/service/<int:service_id>')
def service(service_id):
    data = get_service(service_id)
    return render_template('service.html', service=data["service"], contracts=data["contracts"])

@app.route('/new_service', methods=['GET', 'POST'])
def new_service():
    if request.method == 'POST':
        try:
            name = request.form.get('name', '').strip()
            price = float(request.form.get('price', '0').strip())
            description = request.form.get('description', '').strip()
            if not name or price <= 0:
                flash('Не все обязательные поля заполнены корректно (название, цена).')
                return render_template('new_service.html')
            conn = get_db_connection()
            conn.execute("""
                INSERT INTO services (name, price, description)
                VALUES (?, ?, ?)
            """, (name, price, description))
            conn.commit()
            conn.close()
            flash('Продукт успешно добавлен!')
            return redirect('/services')
        except ValueError:
            flash('Цена должна быть числом.')
            return render_template('new_service.html')
    return render_template('new_service.html')

@app.route('/edit_service/<int:service_id>', methods=['GET', 'POST'])
def edit_service(service_id):
    conn = get_db_connection()
    service = conn.execute('SELECT * FROM services WHERE service_id = ?', (service_id,)).fetchone()
    
    if not service:
        flash("Услуга не найдена.")
        conn.close()
        return redirect('/services')

    if request.method == 'POST':
        name = request.form['name']
        description = request.form['description']
        price = float(request.form['price'])
        conn.execute('UPDATE services SET name = ?, price = ?, description = ? WHERE service_id = ?',
                     (name, price, description, service_id))
        conn.commit()
        conn.close()
        flash('Услуга успешно обновлена.')
        return redirect('/services')
    
    conn.close()
    return render_template('edit_service.html', service=service)


# СОТРУДНИКИ 
def get_employee(item_id):
    conn = get_db_connection()
    employee = conn.execute("""
        SELECT * FROM employees WHERE id_employee = ?
    """, (item_id,)).fetchone()
    if not employee:
        conn.close()
        abort(404)  
    contracts = conn.execute("""
        SELECT 
            contracts.id_contract AS id_contract,
            contracts.number AS number,
            contracts.date AS date,
            contracts.deal_type AS deal_type,
            contracts.start_price AS start_price,
            contracts.discount AS discount,
            contracts.finish_price AS finish_price,
            clients.name AS client_name,
            services.name AS service_name
        FROM contracts
        LEFT JOIN clients ON contracts.client_id = clients.id_client
        LEFT JOIN services ON contracts.service_id = services.service_id
        WHERE contracts.employee_id = ?
    """, (item_id,)).fetchall()
    conn.close()
    employee_data = dict(employee)
    contracts_data = [dict(contract) for contract in contracts]
    return {
        "employee": employee_data,
        "contracts": contracts_data
    }

@app.route('/employee/<int:id_employee>')
def employee(id_employee):
    try:
        data = get_employee(id_employee)
        return render_template(
            'employee.html', 
            employee=data['employee'], 
            contracts=data['contracts']
        )
    except Exception as e:
        print(f"Ошибка: {e}")
        abort(500)

@app.route('/employees')
def employees():
    conn = get_db_connection()
    pos = conn.execute("SELECT * FROM employees").fetchall()
    conn.close()
    return render_template('employees.html', employees=pos)

@app.route('/new_employee', methods=('GET', 'POST'))
def new_employee():
    if request.method == 'POST':
        try:
            name = request.form.get('name', '').strip()
            phone_number = request.form.get('phone_number', '').strip()
            email = request.form.get('email', '').strip()
            position = request.form.get('position', '').strip()
            department = request.form.get('department', '').strip()
        except ValueError:
            flash('Некорректные значения')
            return render_template('new_employee.html')
        if not (name and phone_number and email and position and department):
            flash('Не все поля заполнены')
        else:
            conn = get_db_connection()
            conn.execute("""
                INSERT INTO employees (name, phone_number, email, position, department)
                VALUES (?, ?, ?, ?, ?)
            """, (name, phone_number, email, position, department))
            conn.commit()
            conn.close()
            return redirect('/employees')
    return render_template('new_employee.html')

@app.route('/employee/edit/<int:employee_id>', methods=['GET', 'POST'])
def edit_employee(employee_id):
    conn = get_db_connection()
    employee = conn.execute("SELECT * FROM employees WHERE id_employee = ?", (employee_id,)).fetchone()
    if not employee:
        flash("Сотрудник не найден.")
        return redirect('/employees')
    if request.method == 'POST':
        name = request.form['name'].strip()
        phone_number = request.form['phone_number'].strip()
        email = request.form['email'].strip()
        position = request.form['position'].strip()
        department = request.form['department'].strip()
        if not (name and phone_number and email and position and department):
            flash("Все поля должны быть заполнены.")
        else:
            conn.execute("""
                UPDATE employees
                SET name = ?, phone_number = ?, email = ?, position = ?, department = ?
                WHERE id_employee = ?
            """, (name, phone_number, email, position, department, employee_id))
            conn.commit()
            conn.close()
            flash("Данные сотрудника успешно обновлены.")
            return redirect('/employees')
    conn.close()
    return render_template('edit_employee.html', employee=employee)

# ОТЧЕТЫ 
@app.route('/reports')
def reports():
    conn = get_db_connection()
    reports_data = conn.execute("""
        SELECT
    reports.*,
    employees.name AS employee_name,
    report_types.type_name AS report_type_name
FROM reports
LEFT JOIN employees ON reports.employee_id = employees.id_employee
LEFT JOIN report_types ON reports.report_type_id = report_types.id_type;

    """).fetchall()
    conn.close()
    return render_template('reports.html', reports=reports_data)

@app.route('/new_report', methods=['GET', 'POST'])
def new_report():
    # Получаем данные для селектов
    conn = get_db_connection()
    report_types = conn.execute("SELECT * FROM report_types").fetchall()
    employees = conn.execute("SELECT * FROM employees").fetchall()
    conn.close()

    if request.method == 'POST':
        # Получаем данные из формы
        number = request.form.get('number', '').strip()
        date = request.form.get('date', '').strip()
        report_type_id = request.form.get('report_type')
        description = request.form.get('description', '').strip()
        employee_id = request.form.get('employee_id')

        # Проверяем, что все поля заполнены
        if not (number and date and report_type_id and description and employee_id):
            flash("Все поля должны быть заполнены!", "danger")
            return render_template('new_report.html', report_types=report_types, employees=employees)

        # Преобразуем ID в числа
        try:
            report_type_id = int(report_type_id)
            employee_id = int(employee_id)
        except ValueError:
            flash("Некорректный ID отчета или сотрудника.", "danger")
            return render_template('new_report.html', report_types=report_types, employees=employees)

        # Попытка вставить в базу
        try:
            with get_db_connection() as conn:
                conn.execute("""
                    INSERT INTO reports (number, date, report_type_id, description, employee_id)
                    VALUES (?, ?, ?, ?, ?)
                """, (number, date, report_type_id, description, employee_id))
                conn.commit()
            flash("Отчет успешно добавлен!", "success")
            return redirect(url_for('reports'))
        except sqlite3.IntegrityError:
            flash("Ошибка: отчет с таким номером уже существует.", "danger")
        except Exception as e:
            flash(f"Ошибка при добавлении отчета: {e}", "danger")

    # GET-запрос или ошибка POST — возвращаем форму
    return render_template('new_report.html', report_types=report_types, employees=employees)




@app.route('/generate_report/<int:report_id>', methods=['GET', 'POST'])
def generate_report(report_id):
    """Генерация отчета"""
    conn = get_db_connection()
    report = conn.execute("""
        SELECT 
            reports.*, 
            employees.name AS employee_name
        FROM reports
        LEFT JOIN employees ON reports.employee_id = employees.id_employee
        WHERE id_report = ?
    """, (report_id,)).fetchone()
    conn.close()
    
    if not report:
        flash("Отчет не найден.")
        return redirect('/reports')
    
    report_params = {
        'REPORT_DATE': report['date'],
        'EMPLOYEE_FULLNAME': report['employee_name'],
        'ORDER_DATE_1': '05.04.2024',
        'ORDER_SUM_1': 5000,
        'EMPLOYEE_1': 'Морозов В.А.',
        'EMPLOYEE_2': 'Смеловский Е.А.',
        'EMPLOYEE_3': 'Мирбоев Ф.Д.',
        'ORDER_DATE_2': '07.04.2024',
        'ORDER_SUM_2': 3500,
        'ORDER_DATE_3': '12.04.2024',
        'ORDER_SUM_3': 1500,
        'ORDER_DATE_4': '18.04.2024',
        'ORDER_SUM_4': 12000,
        'TOTAL_ORDERS': 4,
        'TOTAL_SUM': 22000,
        'EMPLOYEE_NAME_1': 'Морозов Василий Алексеевич',
        'EMPLOYEE_1_ORDER_DATE_1': '05.04.2024',
        'EMPLOYEE_1_ORDER_SUM_1': 5000,
        'EMPLOYEE_1_ORDER_DATE_2': '18.04.2024',
        'EMPLOYEE_1_ORDER_SUM_2': 12000,
        'EMPLOYEE_1_TOTAL_ORDERS': 2,
        'EMPLOYEE_1_TOTAL_SUM': 17000,
        'EMPLOYEE_NAME_2': 'Смеловский Егор Андреевич',
        'EMPLOYEE_2_ORDER_DATE_1': '07.04.2024',
        'EMPLOYEE_2_ORDER_SUM_1': 3500,
        'EMPLOYEE_2_ORDER_DATE_2': '25.04.2024',
        'EMPLOYEE_2_ORDER_SUM_2': 7000,
        'EMPLOYEE_2_TOTAL_ORDERS': 2,
        'EMPLOYEE_2_TOTAL_SUM': 10500,
        'EMPLOYEE_NAME_3': 'Мирбоев Фаорис Дониёрович',
        'EMPLOYEE_3_ORDER_DATE_1': '12.04.2024',
        'EMPLOYEE_3_ORDER_SUM_1': 1500,
        'EMPLOYEE_3_ORDER_DATE_2': '28.04.2024',
        'EMPLOYEE_3_ORDER_SUM_2': 10000,
        'EMPLOYEE_3_TOTAL_ORDERS': 2,
        'EMPLOYEE_3_TOTAL_SUM': 11500,
    }
    
    template_path = os.path.join('templates', 'report_template.docx')
    output_path = os.path.join('contracts', f"Отчет_{report['number']}.docx")
    generate_docx(template_path, output_path, report_params)
    
    return send_file(output_path, as_attachment=True)

@app.route('/edit_report/<int:report_id>', methods=['GET', 'POST'])
def edit_report(report_id):
    """Редактирование отчёта"""
    conn = get_db_connection()
    report = conn.execute('SELECT * FROM reports WHERE id_report = ?', (report_id,)).fetchone()
    if not report:
        conn.close()
        flash("Отчёт не найден.")
        return redirect('/reports')
    
    report_types = conn.execute("SELECT * FROM report_types").fetchall()
    employees = conn.execute("SELECT * FROM employees").fetchall()
    conn.close()
    
    if request.method == 'POST':
        try:
            number = request.form['number']
            date = request.form['date']
            report_type = request.form['report_type']
            description = request.form['description']
            employee_id = request.form['employee_id']
        except ValueError:
            flash('Некорректные данные.')
            return render_template('edit_report.html', report=report, report_types=report_types, employees=employees)
        
        if not (number and date and report_type and description and employee_id):
            flash('Все поля должны быть заполнены.')
            return render_template('edit_report.html', report=report, report_types=report_types, employees=employees)

        conn = get_db_connection()
        conn.execute('''
            UPDATE reports
            SET number = ?, date = ?, report_type_id = ?, description = ?, employee_id = ?
            WHERE id_report = ?
        ''', (number, date, report_type, description, employee_id, report_id))
        conn.commit()
        conn.close()
        flash('Отчёт успешно обновлён.')
        return redirect('/reports')
    
    return render_template('edit_report.html', report=report, report_types=report_types, employees=employees)

@app.route('/report/<int:report_id>')
def report(report_id):
    conn = get_db_connection()
    report_data = conn.execute("""
        SELECT 
            reports.*,
            employees.name AS employee_name,
            report_types.type_name AS report_type_name
        FROM reports
        LEFT JOIN employees 
            ON reports.employee_id = employees.id_employee
        LEFT JOIN report_types 
            ON reports.report_type_id = report_types.id_type
        WHERE reports.id_report = ?
    """, (report_id,)).fetchone()
    conn.close()

    if not report_data:
        abort(404)

    return render_template('report.html', report=report_data)


# ОШИБКИ 404  

@app.errorhandler(404)
def page_not_found(e):
    return render_template('404.html'), 404

if __name__ == '__main__':
    app.run(debug=True)

